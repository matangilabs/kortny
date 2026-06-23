"""Unit tests for the doc-editor skill scripts (extract + apply round-trip)."""

from __future__ import annotations

import importlib.util
import types
from pathlib import Path
from typing import Any

import docx

SCRIPTS_DIR = Path(__file__).parent.parent / "kortny/skills/curated/doc-editor/scripts"


def _import_script(name: str) -> types.ModuleType:
    script_path = SCRIPTS_DIR / name
    spec = importlib.util.spec_from_file_location(name.removesuffix(".py"), script_path)
    assert spec is not None, f"Could not create spec for {script_path}"
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def make_styled_docx(path: Path) -> None:
    """Create a small .docx with styled runs for formatting-preservation tests."""
    doc = docx.Document()
    # Paragraph 0: has a bold run and a normal run
    para0 = doc.add_paragraph()
    run1 = para0.add_run("Bold start: ")
    run1.bold = True
    run2 = para0.add_run("normal text here")
    run2.bold = False
    # Paragraph 1: unchanged reference paragraph
    doc.add_paragraph("Second paragraph unchanged")
    # Paragraph 2: target for edit
    doc.add_paragraph("Third paragraph to edit")
    doc.save(str(path))


def test_extract_units_basic(tmp_path: Path) -> None:
    """extract_units returns the correct IDs and sha256s."""
    doc_path = tmp_path / "original.docx"
    make_styled_docx(doc_path)

    extract_mod = _import_script("extract_docx_units.py")
    units: list[dict[str, Any]] = extract_mod.extract_units(doc_path)

    ids = [u["id"] for u in units]
    assert "p-0" in ids
    assert "p-1" in ids
    assert "p-2" in ids

    for u in units:
        if u["id"].startswith("p-"):
            assert len(u["sha256"]) == 16
            assert all(c in "0123456789abcdef" for c in u["sha256"])


def test_apply_patch_changes_target_paragraph(tmp_path: Path) -> None:
    """apply_patch changes p-2 text without touching p-1."""
    doc_path = tmp_path / "original.docx"
    make_styled_docx(doc_path)

    extract_mod = _import_script("extract_docx_units.py")
    apply_mod = _import_script("apply_docx_patch.py")

    units: list[dict[str, Any]] = extract_mod.extract_units(doc_path)

    patch = {"p-2": "Edited third paragraph"}
    result_doc = apply_mod.apply_patch(doc_path, patch, units)

    # Save and reload
    out_path = tmp_path / "revised.docx"
    result_doc.save(str(out_path))
    loaded = docx.Document(str(out_path))

    paragraphs = loaded.paragraphs
    assert paragraphs[2].text == "Edited third paragraph"
    assert paragraphs[1].text == "Second paragraph unchanged"


def test_apply_patch_preserves_bold_on_first_run(tmp_path: Path) -> None:
    """Formatting (bold) on the first run of p-0 is preserved after patching."""
    doc_path = tmp_path / "original.docx"
    make_styled_docx(doc_path)

    extract_mod = _import_script("extract_docx_units.py")
    apply_mod = _import_script("apply_docx_patch.py")

    units: list[dict[str, Any]] = extract_mod.extract_units(doc_path)

    # Patch p-0 (the bold paragraph)
    patch = {"p-0": "Replaced bold paragraph"}
    result_doc = apply_mod.apply_patch(doc_path, patch, units)

    out_path = tmp_path / "revised.docx"
    result_doc.save(str(out_path))
    loaded = docx.Document(str(out_path))

    para = loaded.paragraphs[0]
    assert para.text == "Replaced bold paragraph"
    # First run inherits bold from original run1
    assert para.runs[0].bold is True


def test_extract_sha256_is_deterministic(tmp_path: Path) -> None:
    """Calling extract_units twice returns identical sha256 values."""
    doc_path = tmp_path / "original.docx"
    make_styled_docx(doc_path)

    extract_mod = _import_script("extract_docx_units.py")
    units1: list[dict[str, Any]] = extract_mod.extract_units(doc_path)
    units2: list[dict[str, Any]] = extract_mod.extract_units(doc_path)

    assert [u["sha256"] for u in units1] == [u["sha256"] for u in units2]
