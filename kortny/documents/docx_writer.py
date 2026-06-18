"""Render a Document Studio IR to a Word document (HIG-244 Phase 2).

Same format-agnostic IR as the PDF/PPTX writers; DOCX is the closest fit since
a Word doc is a linear flow of the same editorial blocks. Built with
python-docx (pure-Python, native-editable .docx). Produces a structured,
themed, *editable* document — the right output when the user will keep working
on the text, where a PDF is the right output when it's a finished deliverable.
"""

from __future__ import annotations

import io

from docx import Document as new_document
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from kortny.documents.charts import render_chart_png
from kortny.documents.ir import (
    CTA,
    Callout,
    Chart,
    CoverHeader,
    DocumentSpec,
    Heading,
    Prose,
    PullQuote,
    SectionDivider,
    StatCards,
    Table,
)
from kortny.documents.themes import Theme, resolve_theme


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.lstrip("#"))


def render_docx(spec: DocumentSpec) -> bytes:
    """Render ``spec`` to a .docx file as bytes."""

    theme = resolve_theme(doc_kind=spec.doc_kind, name=spec.theme)
    doc = new_document()
    _set_base_font(doc, theme)
    for block in spec.blocks:
        _render(doc, block, theme)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _set_base_font(doc: Document, theme: Theme) -> None:
    style = doc.styles["Normal"]
    style.font.name = theme.body_font
    style.font.size = Pt(11)
    style.font.color.rgb = _rgb(theme.colors.ink)


def _render(doc: Document, block: object, theme: Theme) -> None:
    if isinstance(block, CoverHeader):
        _cover(doc, block, theme)
    elif isinstance(block, SectionDivider):
        _divider(doc, block, theme)
    elif isinstance(block, Heading):
        _heading(doc, block, theme)
    elif isinstance(block, Prose):
        _prose(doc, block, theme)
    elif isinstance(block, StatCards):
        _stat_cards(doc, block, theme)
    elif isinstance(block, Table):
        _table(doc, block, theme)
    elif isinstance(block, Callout):
        _callout(doc, block, theme)
    elif isinstance(block, PullQuote):
        _pull_quote(doc, block, theme)
    elif isinstance(block, CTA):
        _cta(doc, block, theme)
    elif isinstance(block, Chart):
        _chart(doc, block, theme)


def _chart(doc: Document, block: Chart, theme: Theme) -> None:
    png = render_chart_png(block, theme)
    doc.add_picture(io.BytesIO(png), width=Inches(6.0))
    if block.caption:
        p = doc.add_paragraph()
        _run(
            p,
            block.caption,
            size=9,
            font=theme.mono_font,
            color=theme.colors.muted,
            caps=True,
        )
    doc.add_paragraph()


def _run(
    paragraph: Paragraph,
    text: str,
    *,
    size: int,
    font: str,
    color: str,
    bold: bool = False,
    italic: bool = False,
    caps: bool = False,
) -> None:
    run = paragraph.add_run(text.upper() if caps else text)
    run.font.size = Pt(size)
    run.font.name = font
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = _rgb(color)


def _cover(doc: Document, block: CoverHeader, theme: Theme) -> None:
    c = theme.colors
    if block.eyebrow:
        p = doc.add_paragraph()
        _run(p, block.eyebrow, size=11, font=theme.mono_font, color=c.accent, caps=True)
    p = doc.add_paragraph()
    _run(p, block.title, size=34, font=theme.display_font, color=c.ink, bold=True)
    if block.subtitle:
        p = doc.add_paragraph()
        _run(
            p, block.subtitle, size=15, font=theme.body_font, color=c.muted, italic=True
        )
    if block.meta:
        p = doc.add_paragraph()
        _run(
            p,
            "   ·   ".join(block.meta),
            size=9,
            font=theme.mono_font,
            color=c.muted,
            caps=True,
        )
    doc.add_paragraph()


def _divider(doc: Document, block: SectionDivider, theme: Theme) -> None:
    c = theme.colors
    doc.add_page_break()
    label_bits = [b for b in (block.index, block.label) if b]
    if label_bits:
        p = doc.add_paragraph()
        _run(
            p,
            "  ".join(label_bits),
            size=11,
            font=theme.mono_font,
            color=c.accent,
            caps=True,
        )
    p = doc.add_paragraph()
    _run(p, block.title, size=24, font=theme.display_font, color=c.ink, bold=True)
    if block.subtitle:
        p = doc.add_paragraph()
        _run(
            p, block.subtitle, size=13, font=theme.body_font, color=c.muted, italic=True
        )


def _heading(doc: Document, block: Heading, theme: Theme) -> None:
    p = doc.add_paragraph()
    _run(
        p,
        block.text,
        size=15,
        font=theme.display_font,
        color=theme.colors.ink,
        bold=True,
    )


def _prose(doc: Document, block: Prose, theme: Theme) -> None:
    for para in (s.strip() for s in block.text.split("\n\n")):
        if not para:
            continue
        p = doc.add_paragraph()
        _run(p, para, size=11, font=theme.body_font, color=theme.colors.ink)


def _stat_cards(doc: Document, block: StatCards, theme: Theme) -> None:
    c = theme.colors
    table = doc.add_table(rows=1, cols=len(block.cards))
    for cell, card in zip(table.rows[0].cells, block.cards, strict=True):
        _shade_cell(cell, c.card)
        cell.paragraphs[0].text = ""
        _run(
            cell.paragraphs[0],
            card.value,
            size=20,
            font=theme.display_font,
            color=c.accent,
            bold=True,
        )
        p_label = cell.add_paragraph()
        _run(p_label, card.label, size=11, font=theme.body_font, color=c.ink, bold=True)
        if card.note:
            p_note = cell.add_paragraph()
            _run(p_note, card.note, size=9, font=theme.mono_font, color=c.muted)
    doc.add_paragraph()


def _table(doc: Document, block: Table, theme: Theme) -> None:
    c = theme.colors
    if block.caption:
        p = doc.add_paragraph()
        _run(p, block.caption, size=9, font=theme.mono_font, color=c.muted, caps=True)
    table = doc.add_table(rows=1, cols=len(block.columns))
    table.style = "Table Grid"
    for j, col in enumerate(block.columns):
        cell = table.rows[0].cells[j]
        _shade_cell(cell, c.ink)
        cell.paragraphs[0].text = ""
        _run(
            cell.paragraphs[0],
            col,
            size=10,
            font=theme.mono_font,
            color="#FFFFFF",
            bold=True,
            caps=True,
        )
    for row in block.rows:
        padded = (list(row) + [""] * len(block.columns))[: len(block.columns)]
        cells = table.add_row().cells
        for j, value in enumerate(padded):
            cells[j].paragraphs[0].text = ""
            _run(
                cells[j].paragraphs[0],
                str(value),
                size=10,
                font=theme.body_font,
                color=c.ink,
            )
    doc.add_paragraph()


def _callout(doc: Document, block: Callout, theme: Theme) -> None:
    c = theme.colors
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, theme.panel_tint)
    cell.paragraphs[0].text = ""
    if block.label:
        _run(
            cell.paragraphs[0],
            block.label,
            size=11,
            font=theme.mono_font,
            color=c.accent,
            bold=True,
            caps=True,
        )
        body = cell.add_paragraph()
    else:
        body = cell.paragraphs[0]
    _run(body, block.text, size=12, font=theme.body_font, color=c.ink)
    doc.add_paragraph()


def _pull_quote(doc: Document, block: PullQuote, theme: Theme) -> None:
    c = theme.colors
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    _run(
        p, f'"{block.text}"', size=16, font=theme.display_font, color=c.ink, italic=True
    )
    if block.attribution:
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Pt(18)
        _run(
            pa,
            f"— {block.attribution}",
            size=9,
            font=theme.mono_font,
            color=c.muted,
            caps=True,
        )
    doc.add_paragraph()


def _cta(doc: Document, block: CTA, theme: Theme) -> None:
    c = theme.colors
    if block.text:
        p = doc.add_paragraph()
        _run(p, block.text, size=12, font=theme.body_font, color=c.muted)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(
        p,
        f"➔ {block.label}",
        size=12,
        font=theme.mono_font,
        color=c.accent,
        bold=True,
        caps=True,
    )
    doc.add_paragraph()


def _shade_cell(cell: _Cell, hex_color: str) -> None:
    """Apply a solid background fill to a table cell (no direct python-docx API)."""

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)
