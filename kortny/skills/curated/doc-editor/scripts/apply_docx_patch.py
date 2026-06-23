#!/usr/bin/env python3
"""Apply a text patch to a .docx, preserving run-level formatting."""
import hashlib
import json
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def _sha256_short(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


def _patch_paragraph(para: Paragraph, new_text: str) -> None:
    """Replace paragraph text, preserving the first run's formatting."""
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def apply_patch(doc_path: Path, patch: dict[str, str], units: list[dict]) -> docx.Document:
    """Apply patch to the document, skipping units whose sha256 has changed."""
    doc = docx.Document(str(doc_path))

    sha_lookup: dict[str, str] = {u["id"]: u["sha256"] for u in units}

    para_idx = 0
    tbl_idx = 0

    for child in doc.element.body:
        tag = child.tag
        if tag == qn("w:p"):
            unit_id = f"p-{para_idx}"
            if unit_id in patch:
                para = Paragraph(child, doc)
                current_sha = _sha256_short(para.text)
                if current_sha == sha_lookup.get(unit_id, current_sha):
                    _patch_paragraph(para, patch[unit_id])
                else:
                    print(f"SKIP stale unit {unit_id} (sha changed)", file=sys.stderr)
            para_idx += 1
        elif tag == qn("w:tbl"):
            table = Table(child, doc)
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    unit_id = f"tbl-{tbl_idx}-{r_idx}-{c_idx}"
                    if unit_id in patch:
                        if cell.paragraphs:
                            current_text = "\n".join(p.text for p in cell.paragraphs)
                            current_sha = _sha256_short(current_text)
                            if current_sha == sha_lookup.get(unit_id, current_sha):
                                _patch_paragraph(cell.paragraphs[0], patch[unit_id])
                            else:
                                print(f"SKIP stale unit {unit_id}", file=sys.stderr)
            tbl_idx += 1

    hdr_idx = 0
    ftr_idx = 0
    for section in doc.sections:
        for para in section.header.paragraphs:
            unit_id = f"hdr-{hdr_idx}"
            if unit_id in patch:
                _patch_paragraph(para, patch[unit_id])
            hdr_idx += 1
        for para in section.footer.paragraphs:
            unit_id = f"ftr-{ftr_idx}"
            if unit_id in patch:
                _patch_paragraph(para, patch[unit_id])
            ftr_idx += 1

    return doc


def main() -> None:
    doc_path = Path("/workspace/original.docx")
    patch_path = Path("/workspace/patch.json")
    units_path = Path("/workspace/units.json")
    out_path = Path("/workspace/revised.docx")

    for p in (doc_path, patch_path, units_path):
        if not p.exists():
            print(f"ERROR: {p} not found", file=sys.stderr)
            sys.exit(1)

    patch: dict[str, str] = json.loads(patch_path.read_text())
    units_data = json.loads(units_path.read_text())
    units: list[dict] = units_data.get("units", [])

    doc = apply_patch(doc_path, patch, units)
    doc.save(str(out_path))
    print(f"Saved revised document to {out_path} ({len(patch)} units patched)")


if __name__ == "__main__":
    main()
