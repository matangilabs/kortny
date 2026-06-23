#!/usr/bin/env python3
"""Extract editable text units from a .docx for the doc-editor skill."""
import hashlib
import json
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def _sha256_short(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


def extract_units(doc_path: Path) -> list[dict]:
    doc = docx.Document(str(doc_path))
    units: list[dict] = []
    para_idx = 0
    tbl_idx = 0

    for child in doc.element.body:
        tag = child.tag
        if tag == qn("w:p"):
            para = Paragraph(child, doc)
            text = para.text
            units.append({"id": f"p-{para_idx}", "text": text, "sha256": _sha256_short(text)})
            para_idx += 1
        elif tag == qn("w:tbl"):
            table = Table(child, doc)
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = "\n".join(p.text for p in cell.paragraphs)
                    units.append({"id": f"tbl-{tbl_idx}-{r_idx}-{c_idx}", "text": text, "sha256": _sha256_short(text)})
            tbl_idx += 1

    hdr_idx = 0
    ftr_idx = 0
    for section in doc.sections:
        for para in section.header.paragraphs:
            text = para.text
            units.append({"id": f"hdr-{hdr_idx}", "text": text, "sha256": _sha256_short(text)})
            hdr_idx += 1
        for para in section.footer.paragraphs:
            text = para.text
            units.append({"id": f"ftr-{ftr_idx}", "text": text, "sha256": _sha256_short(text)})
            ftr_idx += 1

    return units


def main() -> None:
    doc_path = Path("/workspace/original.docx")
    out_path = Path("/workspace/units.json")
    if not doc_path.exists():
        print(f"ERROR: {doc_path} not found", file=sys.stderr)
        sys.exit(1)
    units = extract_units(doc_path)
    out_path.write_text(json.dumps({"units": units}, ensure_ascii=False, indent=2))
    print(f"Extracted {len(units)} units to {out_path}")


if __name__ == "__main__":
    main()
